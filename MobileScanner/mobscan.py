#!/usr/bin/env python3
"""
mobscan.py v4 - Unified APK/IPA vulnerability scanner.

Includes:
  * Static analysis (secrets, manifest, WebView bridge, network config, iOS ATS,
    binary hardening, encryption).
  * Evidence-based false-positive VALIDATION ENGINE (weighted signals + baseline).
  * Decompilation (jadx/apktool) and dependency CVE scanning (trivy/grype).
  * Cross-platform Mach-O parsing (macholib).
  * DYNAMIC Frida analysis: real-time SQLi, JS injection, certificate-pinning
    detection/bypass surface, insecure storage, weak crypto, deeplink injection.
  * Report formats: console / text / json / word.

--------------------------------------------------------------------------------
QUICK START
--------------------------------------------------------------------------------
    pip install androguard macholib python-docx frida frida-tools
    # external tools on PATH (optional): jadx, apktool, trivy, grype, otool

    # Static only
    python3 mobscan.py app.apk
    python3 mobscan.py MyApp.ipa --format word --output report.docx

    # Static + dynamic (device/emulator/Corellium with frida-server running)
    python3 mobscan.py app.apk --dynamic --dyn-target com.example.app
    python3 mobscan.py MyApp.ipa --dynamic --dyn-target "AppName" \
            --dyn-host 10.11.1.1 --dyn-duration 45

    # Baseline suppression for CI
    python3 mobscan.py app.apk --baseline baseline.json --auto-suppress
--------------------------------------------------------------------------------
"""

import argparse
import base64
import binascii
import hashlib
import json
import math
import os
import plistlib
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from dataclasses import dataclass, field, asdict
from enum import Enum
from typing import List, Dict, Optional, Tuple


# --------------------------------------------------------------------------- #
# Optional imports (all degrade gracefully)
# --------------------------------------------------------------------------- #
try:
    from androguard.core.apk import APK as AG_APK          # androguard >= 4
    HAVE_ANDROGUARD = True
except Exception:
    try:
        from androguard.core.bytecodes.apk import APK as AG_APK  # androguard < 4
        HAVE_ANDROGUARD = True
    except Exception:
        HAVE_ANDROGUARD = False

try:
    from macholib.MachO import MachO
    from macholib.mach_o import LC_ENCRYPTION_INFO, LC_ENCRYPTION_INFO_64, MH_PIE
    HAVE_MACHOLIB = True
except Exception:
    HAVE_MACHOLIB = False

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

try:
    import frida
    HAVE_FRIDA = True
except Exception:
    HAVE_FRIDA = False


# --------------------------------------------------------------------------- #
# Data model
# --------------------------------------------------------------------------- #
class Severity(str, Enum):
    INFO = "INFO"
    LOW = "LOW"
    MEDIUM = "MEDIUM"
    HIGH = "HIGH"
    CRITICAL = "CRITICAL"


class Confidence(str, Enum):
    LIKELY_FALSE_POSITIVE = "LIKELY_FALSE_POSITIVE"
    NEEDS_REVIEW = "NEEDS_REVIEW"
    CONFIRMED = "CONFIRMED"


@dataclass
class Finding:
    check_id: str
    title: str
    severity: Severity
    confidence: Confidence
    location: str
    evidence: str
    description: str
    validation: str
    remediation: str = ""

    def to_dict(self):
        d = asdict(self)
        d["severity"] = self.severity.value if isinstance(self.severity, Severity) else self.severity
        d["confidence"] = self.confidence.value if isinstance(self.confidence, Confidence) else self.confidence
        return d


SEV_ORDER = {Severity.CRITICAL: 0, Severity.HIGH: 1, Severity.MEDIUM: 2,
             Severity.LOW: 3, Severity.INFO: 4}
SEV_LIST = [Severity.INFO, Severity.LOW, Severity.MEDIUM, Severity.HIGH, Severity.CRITICAL]


def sev_bump(sev: Severity, steps: int) -> Severity:
    i = max(0, min(len(SEV_LIST) - 1, SEV_LIST.index(sev) + steps))
    return SEV_LIST[i]


def _as_sev(v) -> Severity:
    return v if isinstance(v, Severity) else Severity(str(v))


def _conf_val(f: Finding) -> str:
    return f.confidence.value if isinstance(f.confidence, Confidence) else str(f.confidence)


# =========================================================================== #
#  VALIDATION ENGINE
# =========================================================================== #
@dataclass
class Signal:
    weight: float
    reason: str


@dataclass
class FindingContext:
    file_path: str = ""
    package: str = ""
    raw_match: str = ""
    quoted_value: Optional[str] = None
    corroborating_ids: List[str] = field(default_factory=list)
    method_signatures: List[str] = field(default_factory=list)
    declared_only: Optional[bool] = None
    build_variant: Optional[str] = None
    runtime_observed: bool = False
    extra: Dict[str, str] = field(default_factory=dict)


TEST_PATH_RX = re.compile(
    r"(?i)(^|/)(test|androidTest|tests|mock|mocks|sample|samples|examples?|"
    r"fixtures?|debug|__mocks__|spec|stub)s?(/|$)"
)
GENERATED_PATH_RX = re.compile(
    r"(?i)(^|/)(R\.java|BuildConfig\.java|generated|build/|\.gradle/|databinding)"
)

BENIGN_SDKS: Dict[str, Dict] = {
    "com/dynatrace":          {"name": "Dynatrace monitoring SDK", "trust": 0.9},
    "com.dynatrace":          {"name": "Dynatrace monitoring SDK", "trust": 0.9},
    "com/medallia/digital":   {"name": "Medallia Digital survey SDK", "trust": 0.9},
    "com.medallia.digital":   {"name": "Medallia Digital survey SDK", "trust": 0.9},
    "com/google/firebase":    {"name": "Firebase SDK", "trust": 0.85},
    "com/google/android/gms": {"name": "Google Play Services", "trust": 0.85},
    "com/facebook":           {"name": "Facebook SDK", "trust": 0.7},
    "androidx/":              {"name": "AndroidX", "trust": 0.95},
    "kotlin/":                {"name": "Kotlin stdlib", "trust": 0.95},
    "com/squareup/okhttp":    {"name": "OkHttp", "trust": 0.85},
    "com/squareup/retrofit":  {"name": "Retrofit", "trust": 0.85},
    "io/sentry":              {"name": "Sentry SDK", "trust": 0.8},
    "com/appsflyer":          {"name": "AppsFlyer SDK", "trust": 0.7},
    "com/newrelic":           {"name": "New Relic SDK", "trust": 0.8},
    "com/adjust/sdk":         {"name": "Adjust SDK", "trust": 0.7},
    "com/bumptech/glide":     {"name": "Glide", "trust": 0.9},
}


def classify_sdk(path_or_pkg: str) -> Optional[Tuple[str, float]]:
    p = (path_or_pkg or "").replace("\\", "/")
    for marker, meta in BENIGN_SDKS.items():
        if marker in p:
            return meta["name"], meta["trust"]
    return None


PLACEHOLDER_RX = re.compile(
    r"(?i)(example|sample|dummy|test|placeholder|your[_-]?(key|token|secret)|"
    r"xxx+|changeme|123456|foobar|redacted|todo|fixme|lorem|abcdef0*|deadbeef|"
    r"<[^>]+>|\$\{[^}]+\}|%[sd]|\{\{[^}]+\}\})"
)
NON_SECRET_HINT_RX = re.compile(
    r"(?i)(sha256|sha1|md5|checksum|integrity|content[_-]?hash|[0-9a-f]{40}|[0-9a-f]{64})"
)
COMMON_WORDS = {
    "password", "username", "content", "android", "google", "example",
    "application", "development", "production", "configuration", "authorization",
    "authentication", "certificate", "description", "information", "notification",
}


def looks_like_uuid(s: str) -> bool:
    return bool(re.fullmatch(
        r"[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-"
        r"[0-9a-fA-F]{4}-[0-9a-fA-F]{12}", s))


def looks_like_bare_hash(s: str) -> bool:
    return bool(re.fullmatch(r"[0-9a-fA-F]{32}|[0-9a-fA-F]{40}|[0-9a-fA-F]{64}", s))


def dictionary_ratio(s: str) -> float:
    low = s.lower()
    covered = sum(len(w) for w in COMMON_WORDS if w in low)
    return covered / max(1, len(s))


def shannon_entropy(s: str) -> float:
    if not s:
        return 0.0
    freq = {c: s.count(c) for c in set(s)}
    return -sum((n / len(s)) * math.log2(n / len(s)) for n in freq.values())


def try_decode_entropy(s: str) -> Tuple[float, str]:
    best = shannon_entropy(s)
    note = ""
    if re.fullmatch(r"[A-Za-z0-9+/]{16,}={0,2}", s):
        try:
            dec = base64.b64decode(s + "=" * (-len(s) % 4), validate=False)
            txt = dec.decode("utf-8", errors="ignore")
            if txt and sum(c.isprintable() for c in txt) / len(txt) > 0.8:
                de = shannon_entropy(txt)
                if de < best:
                    note = f"decodes (base64) to printable text (entropy {de:.2f})"
                    best = de
        except Exception:
            pass
    if re.fullmatch(r"[0-9a-fA-F]{16,}", s) and len(s) % 2 == 0:
        try:
            dec = binascii.unhexlify(s).decode("utf-8", errors="ignore")
            if dec and sum(c.isprintable() for c in dec) / len(dec) > 0.8:
                de = shannon_entropy(dec)
                if de < best:
                    note = f"decodes (hex) to printable text (entropy {de:.2f})"
                    best = de
        except Exception:
            pass
    return best, note


def validate_secret_format(check_id: str, value: str) -> Optional[Signal]:
    cid = (check_id or "").upper()
    if "AWS_ACCESS_KEY" in cid and re.fullmatch(r"AKIA[0-9A-Z]{16}", value):
        return Signal(+3.0, "Matches exact AWS Access Key ID structure (AKIA + 16).")
    if "GOOGLE_API" in cid and re.fullmatch(r"AIza[0-9A-Za-z\-_]{35}", value):
        return Signal(+2.5, "Matches exact Google API key structure (AIza + 35).")
    if "JWT" in cid:
        parts = value.split(".")
        if len(parts) == 3:
            try:
                hdr = base64.urlsafe_b64decode(parts[0] + "=" * (-len(parts[0]) % 4))
                if b'"alg"' in hdr or b"'alg'" in hdr:
                    return Signal(+2.5, "Decoded JWT header contains 'alg' (structurally valid).")
            except Exception:
                return Signal(-1.0, "JWT-shaped but header does not decode to valid JSON.")
    if "PRIVATE_KEY" in cid:
        return Signal(+4.0, "Literal PEM private-key block present — very high signal.")
    return None


def fingerprint(check_id: str, location: str, evidence: str) -> str:
    return hashlib.sha256(f"{check_id}|{location}|{evidence}".encode()).hexdigest()[:16]


class Baseline:
    def __init__(self, path: Optional[str]):
        self.path = path
        self.suppressed: Dict[str, str] = {}
        if path and os.path.isfile(path):
            try:
                with open(path) as f:
                    self.suppressed = json.load(f).get("suppressed", {})
            except Exception:
                pass

    def is_suppressed(self, fp: str) -> Optional[str]:
        return self.suppressed.get(fp)

    def add(self, fp: str, note: str):
        self.suppressed[fp] = note

    def save(self):
        if not self.path:
            return
        with open(self.path, "w") as f:
            json.dump({"suppressed": self.suppressed}, f, indent=2)


@dataclass
class Verdict:
    confidence: Confidence
    severity_delta: int
    signals: List[Signal]

    @property
    def score(self) -> float:
        return sum(s.weight for s in self.signals)

    def explanation(self) -> str:
        pos = [s.reason for s in self.signals if s.weight > 0]
        neg = [s.reason for s in self.signals if s.weight < 0]
        parts = []
        if pos:
            parts.append("Incriminating: " + "; ".join(pos))
        if neg:
            parts.append("Exonerating: " + "; ".join(neg))
        parts.append(f"[net score {self.score:+.1f}]")
        return " | ".join(parts)


class ValidationEngine:
    T_CONFIRMED = 2.5
    T_REVIEW = -1.0
    T_SUPPRESS = -4.0

    def __init__(self, baseline_path: Optional[str] = None, auto_suppress: bool = False):
        self.baseline = Baseline(baseline_path)
        self.auto_suppress = auto_suppress

    def _location_signals(self, ctx: FindingContext) -> List[Signal]:
        sig = []
        path = ctx.file_path or ""
        if TEST_PATH_RX.search(path):
            sig.append(Signal(-3.0, "In test/mock/sample/debug path — low production impact."))
        if GENERATED_PATH_RX.search(path):
            sig.append(Signal(-2.0, "In generated/build artifact (R.java/BuildConfig/etc.)."))
        sdk = classify_sdk(ctx.package or path)
        if sdk:
            name, trust = sdk
            sig.append(Signal(-3.0 * trust, f"Inside vetted SDK ({name}, trust {trust:.0%})."))
        if ctx.build_variant == "debug":
            sig.append(Signal(-2.0, "Belongs to a debug build variant."))
        elif ctx.build_variant == "release":
            sig.append(Signal(+1.0, "Present in a release build variant."))
        return sig

    def _reachability_signals(self, ctx: FindingContext) -> List[Signal]:
        sig = []
        if ctx.runtime_observed:
            sig.append(Signal(+4.0, "Observed live at runtime via Frida — behavior is real, not inferred."))
        if ctx.declared_only is True:
            sig.append(Signal(-1.5, "Declared but appears unreferenced (dead constant)."))
        elif ctx.declared_only is False:
            sig.append(Signal(+1.0, "Value is referenced/used elsewhere in code."))
        return sig

    def _corroboration_signals(self, ctx: FindingContext) -> List[Signal]:
        sig = []
        corr = set(ctx.corroborating_ids or [])
        risky = {"ANDROID_CLEARTEXT_FLAG", "ANDROID_NSC_CLEARTEXT", "ANDROID_NSC_USER_CA",
                 "IOS_ATS_DISABLED", "DYN_NO_PINNING", "DYN_CLEARTEXT_LOAD"}
        overlap = corr & risky
        if overlap:
            sig.append(Signal(+2.5, f"Corroborated by co-occurring risk finding(s): "
                                    f"{', '.join(sorted(overlap))}."))
        return sig

    def _secret_signals(self, ctx: FindingContext) -> List[Signal]:
        sig = []
        raw = ctx.raw_match or ""
        val = ctx.quoted_value
        if PLACEHOLDER_RX.search(raw):
            sig.append(Signal(-4.0, "Contains placeholder/template markers."))
        if val:
            if looks_like_uuid(val):
                sig.append(Signal(-2.0, "Value is a UUID — usually an identifier, not a secret."))
            if looks_like_bare_hash(val) or NON_SECRET_HINT_RX.search(raw):
                sig.append(Signal(-2.5, "Value looks like a hash/checksum, not a credential."))
            dr = dictionary_ratio(val)
            if dr > 0.5:
                sig.append(Signal(-2.0, f"Mostly dictionary words ({dr:.0%}) — likely config string."))
            ent, note = try_decode_entropy(val)
            if ent < 3.0:
                sig.append(Signal(-2.0, f"Low effective entropy ({ent:.2f})" +
                                        (f"; {note}" if note else "") + " — unlikely live key."))
            elif ent >= 4.0:
                sig.append(Signal(+1.5, f"High entropy ({ent:.2f}) consistent with a real secret."))
            fmt = validate_secret_format(ctx.extra.get("check_id", ""), val)
            if fmt:
                sig.append(fmt)
        return sig

    def _jsbridge_signals(self, ctx: FindingContext) -> List[Signal]:
        sig = []
        methods = ctx.method_signatures or []
        if not methods:
            return sig
        sensitive_rx = re.compile(r"(?i)(token|password|passwd|secret|credential|getFile|"
                                  r"readFile|writeFile|cookie|auth|account|ssn|card|exec|"
                                  r"runtime|shell|download)")
        benign_rx = re.compile(r"(?i)(log|track|analytics|privacy|consent|theme|darkmode|"
                               r"locale|version|ping|heartbeat|count|session|feedback|rating|survey)")
        sensitive = [m for m in methods if sensitive_rx.search(m)]
        benign = [m for m in methods if benign_rx.search(m)]
        if sensitive:
            sig.append(Signal(+3.0, f"Exposed method(s) imply sensitive access: {', '.join(sensitive[:4])}."))
        if benign and not sensitive:
            sig.append(Signal(-2.0, f"All exposed methods look like benign telemetry/UI: {', '.join(benign[:4])}."))
        return sig

    def evaluate(self, finding: Finding, ctx: FindingContext) -> Optional[Finding]:
        ctx.extra.setdefault("check_id", finding.check_id)
        fp = fingerprint(finding.check_id, finding.location, finding.evidence)
        if self.baseline.is_suppressed(fp):
            return None
        signals: List[Signal] = []
        signals += self._location_signals(ctx)
        signals += self._reachability_signals(ctx)
        signals += self._corroboration_signals(ctx)
        cid = finding.check_id.upper()
        if cid.startswith("SECRET_"):
            signals += self._secret_signals(ctx)
        if cid == "ANDROID_JS_BRIDGE":
            signals += self._jsbridge_signals(ctx)
        verdict = self._score_to_verdict(signals)
        finding.severity = sev_bump(_as_sev(finding.severity), verdict.severity_delta)
        finding.confidence = verdict.confidence
        expl = verdict.explanation()
        if expl:
            finding.validation = (finding.validation + "  -- Validation engine: " + expl).strip()
        if self.auto_suppress and verdict.score <= self.T_SUPPRESS:
            self.baseline.add(fp, f"auto-suppressed: net score {verdict.score:+.1f} ({finding.check_id})")
            return None
        return finding

    def _score_to_verdict(self, signals: List[Signal]) -> Verdict:
        score = sum(s.weight for s in signals)
        if score >= self.T_CONFIRMED:
            conf = Confidence.CONFIRMED
            delta = +1 if score >= self.T_CONFIRMED + 2 else 0
        elif score <= self.T_REVIEW:
            conf = Confidence.LIKELY_FALSE_POSITIVE
            delta = -1 if score <= self.T_REVIEW - 1 else 0
        else:
            conf = Confidence.NEEDS_REVIEW
            delta = 0
        return Verdict(confidence=conf, severity_delta=delta, signals=signals)

    def finalize(self):
        self.baseline.save()


# --------------------------------------------------------------------------- #
# Secret detection patterns
# --------------------------------------------------------------------------- #
SECRET_PATTERNS = [
    ("aws_access_key", "AWS Access Key ID", re.compile(r"AKIA[0-9A-Z]{16}"), False),
    ("aws_secret",     "AWS Secret Key",     re.compile(r"(?i)aws.{0,20}?['\"][0-9a-zA-Z/+]{40}['\"]"), True),
    ("google_api",     "Google API Key",     re.compile(r"AIza[0-9A-Za-z\-_]{35}"), False),
    ("private_key",    "Private Key Block",  re.compile(r"-----BEGIN (RSA |EC |DSA |OPENSSH )?PRIVATE KEY-----"), False),
    ("slack_token",    "Slack Token",        re.compile(r"xox[baprs]-[0-9A-Za-z\-]{10,}"), False),
    ("jwt",            "JWT",                re.compile(r"eyJ[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}"), False),
    ("firebase_url",   "Firebase DB URL",    re.compile(r"https://[a-z0-9-]+\.firebaseio\.com"), False),
    ("generic_secret", "Generic secret assignment",
        re.compile(r"(?i)(api[_-]?key|secret|passwd|password|token)\s*[:=]\s*['\"][^'\"]{8,}['\"]"), True),
]


# --------------------------------------------------------------------------- #
# Utility
# --------------------------------------------------------------------------- #
def sha256_of(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def run(cmd: List[str], timeout: int = 300) -> Tuple[int, str, str]:
    try:
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        return p.returncode, p.stdout, p.stderr
    except FileNotFoundError:
        return 127, "", "not found"
    except subprocess.TimeoutExpired:
        return 124, "", "timeout"
    except Exception as e:
        return 1, "", str(e)


def have_tool(name: str) -> bool:
    return shutil.which(name) is not None


def extract_zip(path: str, dest: str):
    with zipfile.ZipFile(path, "r") as z:
        z.extractall(dest)


def walk_files(root: str):
    for dirpath, _, files in os.walk(root):
        for f in files:
            yield os.path.join(dirpath, f)


def read_text(path: str, limit: int = 8_000_000) -> str:
    try:
        with open(path, "rb") as f:
            return f.read(limit).decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _rel_any(path: str, roots: List[str]) -> str:
    for r in roots:
        if r and path.startswith(r):
            return os.path.relpath(path, r)
    return path


def _short(fqcn: str) -> str:
    return fqcn.split(".")[-1] if fqcn else fqcn


QUOTE_RX = re.compile(r"['\"]([^'\"]{6,})['\"]")


def _extract_quoted(evidence: str) -> Optional[str]:
    if not evidence:
        return None
    m = QUOTE_RX.search(evidence)
    if m:
        return m.group(1)
    for rx in (re.compile(r"AKIA[0-9A-Z]{16}"), re.compile(r"AIza[0-9A-Za-z\-_]{35}")):
        m = rx.search(evidence)
        if m:
            return m.group(0)
    return None


METHOD_LIST_RX = re.compile(r"@JavascriptInterface methods?:\s*(.+)$", re.IGNORECASE)


def _methods_from_finding(f: Finding) -> List[str]:
    if f.check_id != "ANDROID_JS_BRIDGE":
        return []
    m = METHOD_LIST_RX.search(f.validation or "")
    if not m:
        return []
    return [s.strip().rstrip(".") for s in m.group(1).split(";") if s.strip()]


def _guess_build_variant(path: str) -> Optional[str]:
    low = (path or "").lower()
    if "/debug/" in low or low.endswith("debug"):
        return "debug"
    if "/release/" in low:
        return "release"
    return None


# --------------------------------------------------------------------------- #
# Decompilation
# --------------------------------------------------------------------------- #
def decompile_apk(apk_path: str, work: str, enabled: bool) -> Dict[str, Optional[str]]:
    result = {"jadx_src": None, "apktool_dir": None}
    if not enabled:
        return result
    if have_tool("jadx"):
        jadx_out = os.path.join(work, "jadx")
        run(["jadx", "-d", jadx_out, "--no-res", "-q", apk_path], timeout=600)
        src = os.path.join(jadx_out, "sources")
        result["jadx_src"] = src if os.path.isdir(src) else (jadx_out if os.path.isdir(jadx_out) else None)
    if have_tool("apktool"):
        apktool_out = os.path.join(work, "apktool")
        run(["apktool", "d", "-f", "-o", apktool_out, apk_path], timeout=600)
        result["apktool_dir"] = apktool_out if os.path.isdir(apktool_out) else None
    return result


# --------------------------------------------------------------------------- #
# Dependency CVE scanning
# --------------------------------------------------------------------------- #
def scan_dependencies(target_dir: str, enabled: bool) -> List[Finding]:
    findings = []
    if not enabled:
        return findings
    if have_tool("trivy"):
        findings += _trivy_scan(target_dir)
    elif have_tool("grype"):
        findings += _grype_scan(target_dir)
    else:
        findings.append(Finding(
            check_id="DEP_SCAN_SKIPPED", title="Dependency CVE scan skipped (no trivy/grype)",
            severity=Severity.INFO, confidence=Confidence.CONFIRMED,
            location=target_dir, evidence="trivy/grype not on PATH",
            description="Could not scan embedded libraries for known CVEs.",
            validation="Install trivy or grype to enable third-party SDK CVE detection.",
        ))
    return findings


def _map_cve_severity(sev: str) -> Severity:
    return {"CRITICAL": Severity.CRITICAL, "HIGH": Severity.HIGH, "MEDIUM": Severity.MEDIUM,
            "MODERATE": Severity.MEDIUM, "LOW": Severity.LOW, "NEGLIGIBLE": Severity.INFO,
            "UNKNOWN": Severity.INFO}.get((sev or "").upper(), Severity.INFO)


def _trivy_scan(target_dir: str) -> List[Finding]:
    findings = []
    code, out, _ = run(["trivy", "fs", "--quiet", "--format", "json",
                        "--scanners", "vuln", target_dir], timeout=600)
    if code != 0 or not out.strip():
        return findings
    try:
        data = json.loads(out)
    except Exception:
        return findings
    for res in data.get("Results", []) or []:
        tgt = res.get("Target", "")
        for v in res.get("Vulnerabilities", []) or []:
            fixed = v.get("FixedVersion", "n/a")
            cve = v.get("VulnerabilityID", "?")
            pkg = v.get("PkgName", "?")
            installed = v.get("InstalledVersion", "?")
            findings.append(Finding(
                check_id=f"CVE_{cve}",
                title=f"Vulnerable dependency: {pkg} {installed} ({cve})",
                severity=_map_cve_severity(v.get("Severity", "UNKNOWN")),
                confidence=Confidence.CONFIRMED if fixed != "n/a" else Confidence.NEEDS_REVIEW,
                location=tgt, evidence=f"{pkg}@{installed} -> {cve}",
                description=v.get("Title", "Known vulnerability in bundled dependency."),
                validation=f"Reported by Trivy. Fixed in: {fixed}. Confirm reachability before rating exploitable.",
                remediation=f"Upgrade {pkg} to {fixed}." if fixed != "n/a" else f"Track upstream fix for {cve}.",
            ))
    return findings


def _grype_scan(target_dir: str) -> List[Finding]:
    findings = []
    code, out, _ = run(["grype", f"dir:{target_dir}", "-q", "-o", "json"], timeout=600)
    if code != 0 or not out.strip():
        return findings
    try:
        data = json.loads(out)
    except Exception:
        return findings
    for match in data.get("matches", []) or []:
        vuln = match.get("vulnerability", {})
        art = match.get("artifact", {})
        cve = vuln.get("id", "?")
        pkg = art.get("name", "?")
        ver = art.get("version", "?")
        fixed = ", ".join(vuln.get("fix", {}).get("versions", []) or []) or "n/a"
        findings.append(Finding(
            check_id=f"CVE_{cve}",
            title=f"Vulnerable dependency: {pkg} {ver} ({cve})",
            severity=_map_cve_severity(vuln.get("severity", "UNKNOWN")),
            confidence=Confidence.NEEDS_REVIEW,
            location=", ".join(l.get("path", "") for l in art.get("locations", []))[:200],
            evidence=f"{pkg}@{ver} -> {cve}",
            description=vuln.get("description", "Known vulnerability in bundled dependency."),
            validation=f"Reported by Grype. Fixed in: {fixed}. Confirm reachability before rating exploitable.",
            remediation=f"Upgrade {pkg} to {fixed}." if fixed != "n/a" else f"Track upstream fix for {cve}.",
        ))
    return findings


# --------------------------------------------------------------------------- #
# Secret scan
# --------------------------------------------------------------------------- #
def scan_secrets(roots: List[str]) -> List[Finding]:
    findings = []
    seen = set()
    text_ext = (".xml", ".json", ".plist", ".txt", ".properties", ".java", ".kt",
                ".js", ".html", ".strings", ".cfg", ".conf", ".smali", ".c", ".m",
                ".mm", ".swift", ".yml", ".yaml")
    for root in roots:
        if not root or not os.path.isdir(root):
            continue
        for path in walk_files(root):
            rel = os.path.relpath(path, root)
            try:
                size = os.path.getsize(path)
            except OSError:
                continue
            if not (path.endswith(text_ext) or size < 2_000_000):
                continue
            content = read_text(path)
            if not content:
                continue
            for sid, name, rx, _need in SECRET_PATTERNS:
                for m in rx.finditer(content):
                    snippet = m.group(0)[:80]
                    key = (sid, snippet)
                    if key in seen:
                        continue
                    seen.add(key)
                    findings.append(Finding(
                        check_id=f"SECRET_{sid.upper()}",
                        title=f"Possible hardcoded secret: {name}",
                        severity=Severity.HIGH, confidence=Confidence.NEEDS_REVIEW,
                        location=rel, evidence=snippet,
                        description="A string matching a known secret/credential format was found.",
                        validation="Pattern matched a secret-like format.",
                        remediation="Move secrets server-side; never ship live credentials. Rotate if confirmed.",
                    ))
    return findings


# --------------------------------------------------------------------------- #
# ANDROID analyzer
# --------------------------------------------------------------------------- #
def analyze_apk(path: str, work: str, decompile: bool) -> Tuple[List[Finding], List[str]]:
    findings: List[Finding] = []
    raw = os.path.join(work, "raw")
    os.makedirs(raw, exist_ok=True)
    extract_zip(path, raw)
    decomp = decompile_apk(path, work, decompile)
    source_roots = [d for d in [decomp["jadx_src"], decomp["apktool_dir"], raw] if d]
    findings += check_android_manifest(path, raw, decomp["apktool_dir"])
    findings += check_android_webview(source_roots)
    findings += check_android_network(source_roots)
    if any(f.endswith(".so") for f in walk_files(raw)):
        findings.append(Finding(
            check_id="ANDROID_NATIVE_LIBS", title="Native libraries present",
            severity=Severity.INFO, confidence=Confidence.CONFIRMED,
            location="lib/", evidence="*.so found",
            description="App ships native code; consider separate binary hardening review.",
            validation="Informational.",
        ))
    return findings, source_roots


def check_android_manifest(apk_path: str, raw: str, apktool_dir: Optional[str]) -> List[Finding]:
    if HAVE_ANDROGUARD:
        try:
            return _androguard_manifest(apk_path)
        except Exception:
            pass
    mtext = ""
    if apktool_dir:
        mtext = read_text(os.path.join(apktool_dir, "AndroidManifest.xml"))
    if not mtext:
        mtext = read_text(os.path.join(raw, "AndroidManifest.xml"))
    findings = []
    for needle, cid, title, sev, desc, rem in [
        ('android:debuggable="true"', "ANDROID_DEBUGGABLE", "Application is debuggable",
         Severity.HIGH, "Debuggable release lets attackers attach a debugger.",
         "Set android:debuggable=false in release builds."),
        ('android:allowBackup="true"', "ANDROID_ALLOWBACKUP", "Backups allowed (allowBackup=true)",
         Severity.MEDIUM, "App data extractable via adb backup on some devices.",
         "Set android:allowBackup=false unless required."),
        ('android:usesCleartextTraffic="true"', "ANDROID_CLEARTEXT_FLAG", "Cleartext traffic permitted",
         Severity.MEDIUM, "App may send unencrypted HTTP.", "Disable cleartext; enforce HTTPS."),
    ]:
        if needle in mtext:
            findings.append(Finding(
                check_id=cid, title=title, severity=sev, confidence=Confidence.CONFIRMED,
                location="AndroidManifest.xml", evidence=needle, description=desc,
                validation="Parsed from decoded manifest. Confirm it applies to the release build.",
                remediation=rem,
            ))
    return findings


def _androguard_manifest(apk_path: str) -> List[Finding]:
    findings = []
    a = AG_APK(apk_path)
    for attr, cid, title, sev, desc, rem in [
        ("debuggable", "ANDROID_DEBUGGABLE", "Application is debuggable", Severity.HIGH,
         "Debuggable release lets attackers attach a debugger.",
         "Set android:debuggable=false in release builds."),
        ("allowBackup", "ANDROID_ALLOWBACKUP", "Backups allowed (allowBackup=true)", Severity.MEDIUM,
         "App data can be extracted via adb backup on some devices.",
         "Set android:allowBackup=false unless required."),
        ("usesCleartextTraffic", "ANDROID_CLEARTEXT_FLAG", "Cleartext traffic permitted", Severity.MEDIUM,
         "App is allowed to send/receive unencrypted HTTP traffic.",
         "Disable cleartext traffic; enforce HTTPS."),
    ]:
        if a.get_attribute_value("application", attr) == "true":
            findings.append(Finding(
                check_id=cid, title=title, severity=sev, confidence=Confidence.CONFIRMED,
                location="AndroidManifest.xml", evidence=f"android:{attr}=true",
                description=desc, validation="Authoritatively parsed via androguard.", remediation=rem,
            ))
    for kind, getter in (("activity", a.get_activities), ("service", a.get_services),
                         ("receiver", a.get_receivers), ("provider", a.get_providers)):
        for comp in getter():
            if a.get_attribute_value(kind, "exported", name=comp) == "true":
                permission = a.get_attribute_value(kind, "permission", name=comp)
                if permission:
                    sev, conf = Severity.LOW, Confidence.NEEDS_REVIEW
                    validation = (f"Exported {kind} '{_short(comp)}' guarded by permission "
                                  f"'{permission}'. Verify its protectionLevel is signature/dangerous.")
                else:
                    sev, conf = Severity.MEDIUM, Confidence.NEEDS_REVIEW
                    validation = (f"Exported {kind} '{_short(comp)}' has NO permission guard. "
                                  "Verify it validates all incoming intent data.")
                findings.append(Finding(
                    check_id=f"ANDROID_EXPORTED_{kind.upper()}",
                    title=f"Exported {kind}: {_short(comp)}", severity=sev, confidence=conf,
                    location="AndroidManifest.xml",
                    evidence=f"exported=true permission={permission or 'NONE'}",
                    description=f"{kind.capitalize()} is exported and invokable by other apps.",
                    validation=validation,
                    remediation="Add a signature-level permission, validate intent input, or set exported=false.",
                ))
    return findings


def check_android_webview(source_roots: List[str]) -> List[Finding]:
    findings = []
    has_js_bridge = False
    bridge_locations = set()
    loads_remote_http = False
    ssl_bypass = False
    exposed_methods: List[Tuple[str, str]] = []
    js_bridge_rx = re.compile(r"addJavascriptInterface")
    loadurl_http_rx = re.compile(r"loadUrl$$\s*[\"']http://")
    ssl_err_rx = re.compile(r"onReceivedSslError")
    ssl_proceed_rx = re.compile(r"\.proceed$$")
    js_iface_rx = re.compile(r"@JavascriptInterface\s*(?:public\s+)?[\w<>$$$$]+\s+(\w+)\s*$$([^)]*)$$")
    scanned = set()
    for root in source_roots:
        for path in walk_files(root):
            if not path.endswith((".java", ".kt", ".smali")):
                continue
            if os.path.getsize(path) > 8_000_000:
                continue
            content = read_text(path)
            if not content:
                continue
            if js_bridge_rx.search(content):
                has_js_bridge = True
                bridge_locations.add(_rel_any(path, source_roots))
            if loadurl_http_rx.search(content):
                loads_remote_http = True
            if ssl_err_rx.search(content) and ssl_proceed_rx.search(content):
                ssl_bypass = True
            for m in js_iface_rx.finditer(content):
                sig = f"{m.group(1)}({m.group(2).strip()})"
                relf = _rel_any(path, source_roots)
                if (relf, sig) not in scanned:
                    scanned.add((relf, sig))
                    exposed_methods.append((relf, sig))
    if not has_js_bridge and not exposed_methods:
        return findings
    if loads_remote_http or ssl_bypass:
        sev, conf = Severity.HIGH, Confidence.CONFIRMED
        validation = ("Corroborated: JS bridge PLUS " +
                      ("cleartext HTTP loadUrl " if loads_remote_http else "") +
                      ("and SSL-error bypass " if ssl_bypass else "") +
                      "-> untrusted content can reach native code.")
    else:
        sev, conf = Severity.MEDIUM, Confidence.NEEDS_REVIEW
        validation = ("JS bridge present; no HTTP/SSL-bypass evidence statically. "
                      "Confirm WebView loads only trusted HTTPS/local content.")
    if exposed_methods:
        validation += " @JavascriptInterface methods: " + "; ".join(s for _, s in exposed_methods[:15]) + "."
    findings.append(Finding(
        check_id="ANDROID_JS_BRIDGE",
        title="WebView JavaScript bridge (addJavascriptInterface)",
        severity=sev, confidence=conf,
        location="; ".join(sorted(bridge_locations))[:300] or "source",
        evidence="addJavascriptInterface" +
                 (" + loadUrl(http://)" if loads_remote_http else "") +
                 (" + onReceivedSslError.proceed()" if ssl_bypass else "") +
                 (f" + {len(exposed_methods)} exposed method(s)" if exposed_methods else ""),
        description="Native code is exposed to JavaScript via a WebView bridge.",
        validation=validation,
        remediation="Load only trusted HTTPS/local content, minimize exposed methods, "
                    "restrict navigation with a WebViewClient whitelist, never bypass SSL errors.",
    ))
    return findings


def check_android_network(source_roots: List[str]) -> List[Finding]:
    findings = []
    for root in source_roots:
        for path in walk_files(root):
            if path.endswith("network_security_config.xml"):
                content = read_text(path)
                if 'cleartextTrafficPermitted="true"' in content:
                    findings.append(Finding(
                        check_id="ANDROID_NSC_CLEARTEXT",
                        title="Network security config permits cleartext",
                        severity=Severity.MEDIUM, confidence=Confidence.CONFIRMED,
                        location=_rel_any(path, source_roots),
                        evidence='cleartextTrafficPermitted="true"',
                        description="Config explicitly allows cleartext HTTP.",
                        validation="Confirm which domains this applies to; scoped exceptions may be OK.",
                        remediation="Restrict to specific domains or remove; enforce HTTPS.",
                    ))
                if "<trust-anchors>" in content and "user" in content:
                    findings.append(Finding(
                        check_id="ANDROID_NSC_USER_CA",
                        title="Network config trusts user-added CAs",
                        severity=Severity.MEDIUM, confidence=Confidence.NEEDS_REVIEW,
                        location=_rel_any(path, source_roots),
                        evidence="user trust-anchor",
                        description="App trusts user-installed certificates, easing MitM.",
                        validation="Acceptable for debug builds only; confirm not in release.",
                        remediation="Trust only system CAs in release; pin critical endpoints.",
                    ))
    return findings


# --------------------------------------------------------------------------- #
# iOS analyzer
# --------------------------------------------------------------------------- #
def analyze_ipa(path: str, work: str) -> Tuple[List[Finding], List[str]]:
    findings: List[Finding] = []
    raw = os.path.join(work, "raw")
    os.makedirs(raw, exist_ok=True)
    extract_zip(path, raw)
    app_dir = find_app_dir(raw)
    if not app_dir:
        findings.append(Finding(
            check_id="IPA_STRUCTURE", title="No .app bundle found",
            severity=Severity.INFO, confidence=Confidence.CONFIRMED,
            location="Payload/", evidence="missing .app",
            description="Could not locate the .app bundle inside Payload/.",
            validation="Verify this is a valid IPA.",
        ))
        return findings, [raw]
    info_plist = os.path.join(app_dir, "Info.plist")
    plist = load_plist(info_plist)
    findings += check_ios_ats(plist, os.path.relpath(info_plist, raw))
    findings += check_ios_urlschemes(plist, os.path.relpath(info_plist, raw))
    binary = find_macho_binary(app_dir)
    if binary:
        enc_finding, encrypted = check_ios_encryption(binary, raw)
        if enc_finding:
            findings.append(enc_finding)
        findings += check_ios_binary_hardening(binary, raw)
        if encrypted:
            findings.append(Finding(
                check_id="IOS_STATIC_LIMITED", title="Static analysis limited (encrypted binary)",
                severity=Severity.INFO, confidence=Confidence.CONFIRMED,
                location=os.path.relpath(binary, raw), evidence="cryptid=1",
                description="Secret/string scanning unreliable while encrypted.",
                validation="Run frida-ios-dump on a jailbroken/Corellium device, then re-scan.",
            ))
    return findings, [raw]


def find_app_dir(root: str) -> Optional[str]:
    payload = os.path.join(root, "Payload")
    base = payload if os.path.isdir(payload) else root
    if os.path.isdir(base):
        for entry in os.listdir(base):
            if entry.endswith(".app"):
                return os.path.join(base, entry)
    return None


def load_plist(path: str) -> dict:
    try:
        with open(path, "rb") as f:
            return plistlib.load(f)
    except Exception:
        if have_tool("plutil"):
            tmp = path + ".xml"
            run(["plutil", "-convert", "xml1", "-o", tmp, path])
            try:
                with open(tmp, "rb") as f:
                    return plistlib.load(f)
            except Exception:
                return {}
        return {}


def check_ios_ats(plist: dict, loc: str) -> List[Finding]:
    findings = []
    ats = plist.get("NSAppTransportSecurity", {}) or {}
    if ats.get("NSAllowsArbitraryLoads") is True:
        if bool(ats.get("NSExceptionDomains")):
            sev, conf = Severity.LOW, Confidence.NEEDS_REVIEW
            validation = ("Arbitrary loads enabled but exception domains defined — on iOS 10+ "
                          "scoped exceptions may override. Review the domain list.")
        else:
            sev, conf = Severity.HIGH, Confidence.CONFIRMED
            validation = "Unscoped NSAllowsArbitraryLoads=true disables ATS app-wide."
        findings.append(Finding(
            check_id="IOS_ATS_DISABLED", title="App Transport Security weakened",
            severity=sev, confidence=conf, location=loc, evidence="NSAllowsArbitraryLoads=true",
            description="ATS disabled/weakened, permitting insecure HTTP.",
            validation=validation,
            remediation="Remove NSAllowsArbitraryLoads; use narrowly scoped exception domains only.",
        ))
    return findings


def check_ios_urlschemes(plist: dict, loc: str) -> List[Finding]:
    findings = []
    schemes = []
    for entry in plist.get("CFBundleURLTypes", []) or []:
        schemes += entry.get("CFBundleURLSchemes", []) or []
    if schemes:
        findings.append(Finding(
            check_id="IOS_CUSTOM_URL_SCHEME", title="Custom URL scheme(s) registered",
            severity=Severity.LOW, confidence=Confidence.NEEDS_REVIEW,
            location=loc, evidence=", ".join(schemes)[:120],
            description="App registers custom URL schemes invokable by other apps.",
            validation="Not a vuln alone. Verify deep-link handlers validate input.",
            remediation="Validate deep-link input; prefer Universal Links for sensitive flows.",
        ))
    return findings


def find_macho_binary(app_dir: str) -> Optional[str]:
    plist = load_plist(os.path.join(app_dir, "Info.plist"))
    exe = plist.get("CFBundleExecutable")
    if exe:
        cand = os.path.join(app_dir, exe)
        if os.path.isfile(cand):
            return cand
    best, best_size = None, 0
    for f in os.listdir(app_dir):
        full = os.path.join(app_dir, f)
        if os.path.isfile(full) and "." not in f and os.path.getsize(full) > best_size:
            best, best_size = full, os.path.getsize(full)
    return best


def check_ios_encryption(binary: str, root: str) -> Tuple[Optional[Finding], bool]:
    rel = os.path.relpath(binary, root)
    cryptid = _get_cryptid(binary)
    if cryptid is None:
        return None, False
    if cryptid == 0:
        return Finding(
            check_id="IOS_NOT_ENCRYPTED", title="Binary not FairPlay-encrypted (cryptid=0)",
            severity=Severity.INFO, confidence=Confidence.CONFIRMED,
            location=rel, evidence="cryptid 0",
            description="Binary is decrypted — static analysis reliable.",
            validation="Good for analysis.",
        ), False
    return Finding(
        check_id="IOS_ENCRYPTED", title="Binary FairPlay-encrypted (cryptid=1)",
        severity=Severity.INFO, confidence=Confidence.CONFIRMED,
        location=rel, evidence="cryptid 1",
        description="Encrypted binary — static findings incomplete.",
        validation="Decrypt with frida-ios-dump, then re-scan.",
    ), True


def _get_cryptid(binary: str) -> Optional[int]:
    if HAVE_MACHOLIB:
        try:
            m = MachO(binary)
            for header in m.headers:
                for lc, cmd, _data in header.commands:
                    if lc.cmd in (LC_ENCRYPTION_INFO, LC_ENCRYPTION_INFO_64):
                        return getattr(cmd, "cryptid", None)
            return 0
        except Exception:
            pass
    if have_tool("otool"):
        _, out, _ = run(["otool", "-l", binary])
        m = re.search(r"cryptid\s+(\d+)", out)
        if m:
            return int(m.group(1))
    return None


def check_ios_binary_hardening(binary: str, root: str) -> List[Finding]:
    findings = []
    rel = os.path.relpath(binary, root)
    if _has_pie(binary) is False:
        findings.append(Finding(
            check_id="IOS_NO_PIE", title="Binary not compiled with PIE",
            severity=Severity.MEDIUM, confidence=Confidence.CONFIRMED,
            location=rel, evidence="MH_PIE flag absent",
            description="No Position Independent Executable — weakens ASLR.",
            validation="Parsed from Mach-O header flags.",
            remediation="Compile with -fPIE -pie.",
        ))
    symbols = _read_strings(binary)
    if symbols is not None:
        if "__stack_chk_guard" not in symbols and "__stack_chk_fail" not in symbols:
            findings.append(Finding(
                check_id="IOS_NO_STACK_CANARY", title="No stack canary detected",
                severity=Severity.MEDIUM, confidence=Confidence.NEEDS_REVIEW,
                location=rel, evidence="__stack_chk_* not found",
                description="Stack-smashing protection may be absent.",
                validation="Heuristic (symbol scan); may be stripped.",
                remediation="Build with -fstack-protector-all.",
            ))
        if "_objc_release" not in symbols and "_objc_retainAutorelease" not in symbols:
            findings.append(Finding(
                check_id="IOS_NO_ARC", title="ARC possibly not enabled",
                severity=Severity.LOW, confidence=Confidence.NEEDS_REVIEW,
                location=rel, evidence="ARC symbols not found",
                description="Automatic Reference Counting may be disabled.",
                validation="Heuristic based on symbol presence.",
                remediation="Enable ARC.",
            ))
    return findings


def _has_pie(binary: str) -> Optional[bool]:
    if HAVE_MACHOLIB:
        try:
            m = MachO(binary)
            for header in m.headers:
                return bool(header.header.flags & MH_PIE)
        except Exception:
            pass
    if have_tool("otool"):
        _, out, _ = run(["otool", "-hv", binary])
        return "PIE" in out
    return None


def _read_strings(binary: str, limit: int = 30_000_000) -> Optional[str]:
    try:
        with open(binary, "rb") as f:
            data = f.read(limit)
    except Exception:
        return None
    return "".join(chr(b) if 32 <= b < 127 else "\n" for b in data)


# =========================================================================== #
#  DYNAMIC ANALYSIS (Frida)
# =========================================================================== #
ANDROID_HOOKS = r"""
'use strict';
function F(id, sev, title, ev, det){ send({t:'finding',check_id:id,severity:sev,title:title,evidence:ev,detail:det}); }
Java.perform(function () {
    try {
        var DB = Java.use('android.database.sqlite.SQLiteDatabase');
        ['rawQuery','execSQL'].forEach(function(m){
            DB[m].overloads.forEach(function(ov){
                ov.implementation = function(){
                    var sql = arguments[0]?arguments[0].toString():'';
                    if ((/'.*\+.*'/.test(sql)) || (/=\s*'[^?]/.test(sql) && sql.indexOf('?')===-1)) {
                        F('DYN_SQLI','HIGH','Potential SQL injection (dynamic string query)',sql.substring(0,200),
                          'Query executed without parameterized placeholders (?). Input reaching this is injectable.');
                    }
                    return ov.apply(this,arguments);
                };
            });
        });
    } catch(e){}
    try {
        var WV = Java.use('android.webkit.WebView');
        ['loadUrl','evaluateJavascript','loadData','loadDataWithBaseURL'].forEach(function(m){
            if(!WV[m]) return;
            WV[m].overloads.forEach(function(ov){
                ov.implementation = function(){
                    var a = arguments[0]?arguments[0].toString():'';
                    if(m==='loadUrl' && a.indexOf('javascript:')===0)
                        F('DYN_JS_INJECT','MEDIUM','Runtime javascript: URL loaded into WebView',a.substring(0,200),
                          'javascript: scheme executed at runtime -- injection vector if arg is tainted.');
                    if(a.indexOf('http://')===0)
                        F('DYN_CLEARTEXT_LOAD','HIGH','WebView loaded cleartext HTTP at runtime',a.substring(0,200),
                          'Live HTTP load enables MitM JS injection into the WebView bridge.');
                    return ov.apply(this,arguments);
                };
            });
        });
    } catch(e){}
    var pinning=false;
    try {
        var CP = Java.use('okhttp3.CertificatePinner');
        CP.check.overload('java.lang.String','java.util.List').implementation = function(h,c){
            pinning=true;
            F('DYN_PINNING_PRESENT','INFO','OkHttp certificate pinning active','okhttp3.CertificatePinner.check('+h+')',
              'Pinning enforced for '+h+'. Bypassable via this same hook.');
            return this.check(h,c);
        };
    } catch(e){}
    try {
        var TM = Java.use('com.android.org.conscrypt.TrustManagerImpl');
        TM.checkTrustedRecursive.implementation = function(){
            pinning=true;
            F('DYN_PINNING_CONSCRYPT','INFO','Platform pinning (Conscrypt) invoked','TrustManagerImpl.checkTrustedRecursive',
              'NSC pinning in force but bypassable by hooking this method.');
            return this.checkTrustedRecursive.apply(this,arguments);
        };
    } catch(e){}
    setTimeout(function(){
        if(!pinning)
            F('DYN_NO_PINNING','MEDIUM','No certificate pinning observed at runtime','no pinning APIs invoked',
              'App made TLS connections without observed pinning -- MitM-vulnerable with a trusted CA. Confirm via proxy.');
    },15000);
    try {
        var Ed = Java.use('android.content.SharedPreferences$Editor');
        Ed.putString.implementation = function(k,v){
            var kk=k?k.toString():'', vv=v?v.toString():'';
            if(/(?:token|password|secret|pin|ssn|card|key)/i.test(kk) || /eyJ[A-Za-z0-9_-]{10,}\./.test(vv))
                F('DYN_INSECURE_STORAGE','HIGH','Sensitive value written to SharedPreferences (plaintext)',
                  kk+'='+vv.substring(0,40),'Sensitive data stored unencrypted. Use EncryptedSharedPreferences/Keystore.');
            return this.putString(k,v);
        };
    } catch(e){}
    try {
        var C = Java.use('javax.crypto.Cipher');
        C.getInstance.overload('java.lang.String').implementation = function(t){
            var tt=t.toString();
            if(/DES|RC4|ECB|MD5/i.test(tt))
                F('DYN_WEAK_CRYPTO','MEDIUM','Weak cryptographic transformation used',tt,
                  'Insecure algorithm/mode ('+tt+'). Use AES/GCM with random IV.');
            return this.getInstance(t);
        };
    } catch(e){}
    try {
        var In = Java.use('android.content.Intent');
        In.getData.implementation = function(){
            var u=this.getData();
            if(u){ var s=u.toString();
                if(/['"<>]|javascript:|file:\/\//i.test(s))
                    F('DYN_DEEPLINK_INJECT','MEDIUM','Suspicious deep-link/intent data received',s.substring(0,200),
                      'Incoming URI contains injection-adjacent chars/schemes. Verify handler sanitizes this.');
            }
            return u;
        };
    } catch(e){}
    send({t:'status',msg:'Android hooks installed'});
});
"""

IOS_HOOKS = r"""
'use strict';
function F(id,sev,title,ev,det){ send({t:'finding',check_id:id,severity:sev,title:title,evidence:ev,detail:det}); }
if (ObjC.available) {
    try {
        var prep = Module.findExportByName(null,'sqlite3_prepare_v2');
        if(prep) Interceptor.attach(prep,{ onEnter:function(a){
            var sql=a[1].readUtf8String();
            if(sql && (/'.*%@.*'|'.*\+.*'/.test(sql)) && sql.indexOf('?')===-1)
                F('DYN_SQLI','HIGH','Potential SQL injection (unparameterized query)',sql.substring(0,200),
                  'Query built via string interpolation without bound parameters (?).');
        }});
    } catch(e){}
    var pinning=false;
    try {
        if(ObjC.classes.AFSecurityPolicy){
            Interceptor.attach(ObjC.classes.AFSecurityPolicy['- evaluateServerTrust:forDomain:'].implementation,{
                onEnter:function(){ pinning=true;
                    F('DYN_PINNING_PRESENT','INFO','AFNetworking certificate pinning active','AFSecurityPolicy.evaluateServerTrust',
                      'Pinning enforced via AFNetworking -- bypassable by hooking this method.'); }
            });
        }
        if(ObjC.classes.TSKPinningValidator){ pinning=true;
            F('DYN_PINNING_PRESENT','INFO','TrustKit pinning present','TSKPinningValidator',
              'TrustKit enforces pinning -- hookable/bypassable at runtime.'); }
    } catch(e){}
    try {
        var se = Module.findExportByName('Security','SecTrustEvaluateWithError');
        if(se) Interceptor.attach(se,{ onLeave:function(r){ pinning=true;
            F('DYN_PINNING_SECTRUST','INFO','SecTrustEvaluate invoked (custom trust check)','Security.SecTrustEvaluateWithError',
              'Custom trust evaluation present. Bypassable by forcing success.'); }});
    } catch(e){}
    setTimeout(function(){
        if(!pinning)
            F('DYN_NO_PINNING','MEDIUM','No certificate pinning observed at runtime','no pinning APIs invoked',
              'No pinning APIs fired during TLS. Confirm via proxy; if traffic flows through a trusted proxy CA, MitM-vulnerable.');
    },15000);
    try {
        var UD = ObjC.classes.NSUserDefaults;
        Interceptor.attach(UD['- setObject:forKey:'].implementation,{ onEnter:function(a){
            var k=new ObjC.Object(a[3]).toString(), v=new ObjC.Object(a[2]).toString();
            if(/token|password|secret|pin|ssn|card|key/i.test(k))
                F('DYN_INSECURE_STORAGE','HIGH','Sensitive value written to NSUserDefaults (plaintext)',
                  k+'='+v.substring(0,40),'NSUserDefaults is unencrypted plist storage. Use Keychain for secrets.');
        }});
    } catch(e){}
    try {
        var cc = Module.findExportByName('libcommonCrypto.dylib','CCCrypt') ||
                 Module.findExportByName(null,'CCCrypt');
        if(cc) Interceptor.attach(cc,{ onEnter:function(a){
            var alg=a[1].toInt32(), opts=a[2].toInt32();
            if(alg===1||alg===4||(opts&2))
                F('DYN_WEAK_CRYPTO','MEDIUM','Weak crypto via CommonCrypto','CCCrypt alg='+alg+' opts='+opts,
                  'DES/RC4 or ECB mode in use. Switch to AES-GCM with random IV.');
        }});
    } catch(e){}
    send({t:'status',msg:'iOS hooks installed'});
}
"""


def run_dynamic(platform: str, target: str, host: Optional[str] = None,
                spawn: bool = True, duration: int = 30) -> List[dict]:
    if not HAVE_FRIDA:
        print("[DYN] error: frida not installed. Run: pip install frida frida-tools", file=sys.stderr)
        return []
    collected: List[dict] = []

    def on_message(message, data):
        if message.get("type") == "send":
            p = message.get("payload", {})
            if p.get("t") == "finding":
                collected.append(p)
                print(f"[DYN][{p['severity']}] {p['title']} :: {p['evidence'][:80]}")
            elif p.get("t") == "status":
                print(f"[DYN] {p.get('msg')}")
        elif message.get("type") == "error":
            print(f"[DYN][frida-error] {message.get('description')}", file=sys.stderr)

    try:
        device = (frida.get_device_manager().add_remote_device(host) if host
                  else frida.get_usb_device(timeout=10))
    except Exception as e:
        print(f"[DYN] error: could not get device: {e}", file=sys.stderr)
        return []

    script_src = ANDROID_HOOKS if platform == "android" else IOS_HOOKS
    try:
        if spawn:
            pid = device.spawn([target])
            session = device.attach(pid)
            script = session.create_script(script_src)
            script.on("message", on_message)
            script.load()
            device.resume(pid)
        else:
            session = device.attach(target)
            script = session.create_script(script_src)
            script.on("message", on_message)
            script.load()
    except Exception as e:
        print(f"[DYN] error: could not attach/spawn '{target}': {e}", file=sys.stderr)
        return collected

    print(f"[DYN] Instrumenting for {duration}s -- EXERCISE THE APP NOW "
          f"(log in, navigate, submit forms) to trigger code paths.")
    try:
        time.sleep(duration)
    except KeyboardInterrupt:
        pass
    try:
        session.detach()
    except Exception:
        pass
    return collected


def dynamic_to_findings(dyn: List[dict]) -> List[Finding]:
    out = []
    for d in dyn:
        out.append(Finding(
            check_id=d["check_id"], title=d["title"],
            severity=Severity(d["severity"]), confidence=Confidence.CONFIRMED,
            location="[runtime]", evidence=d["evidence"], description=d["detail"],
            validation="Observed dynamically via Frida instrumentation -- live, runtime-confirmed behavior.",
            remediation="",
        ))
    return out


# --------------------------------------------------------------------------- #
# Reporting
# --------------------------------------------------------------------------- #
def sort_findings(findings: List[Finding]) -> List[Finding]:
    return sorted(findings, key=lambda f: (SEV_ORDER[_as_sev(f.severity)], _conf_val(f)))


def build_summary(findings: List[Finding]) -> Dict[str, int]:
    counts = {s.value: 0 for s in Severity}
    for f in findings:
        counts[_as_sev(f.severity).value] += 1
    return counts


def report_console(meta: dict, findings: List[Finding]):
    findings = sort_findings(findings)
    print("=" * 78)
    print(" Mobile App Scan Report")
    print(f" Target : {meta['target']}")
    print(f" Type   : {meta['type']}")
    print(f" SHA256 : {meta['sha256']}")
    print(f" Tools  : {meta['tooling']}")
    print(f" Findings: {len(findings)}")
    print("=" * 78)
    counts = build_summary(findings)
    print(" Severity: " + ", ".join(f"{k}={v}" for k, v in counts.items() if v))
    fp = sum(1 for f in findings if _conf_val(f) == Confidence.LIKELY_FALSE_POSITIVE.value)
    print(f" Likely false positives (auto-downgraded): {fp}")
    print("-" * 78)
    for f in findings:
        flag = " [LIKELY FP]" if _conf_val(f) == Confidence.LIKELY_FALSE_POSITIVE.value else ""
        print(f"\n[{_as_sev(f.severity).value}]{flag} {f.title}  ({f.check_id})")
        print(f"  Confidence : {_conf_val(f)}")
        print(f"  Location   : {f.location}")
        print(f"  Evidence   : {f.evidence}")
        print(f"  Detail     : {f.description}")
        print(f"  Validation : {f.validation}")
        if f.remediation:
            print(f"  Fix        : {f.remediation}")
    print("\n" + "=" * 78)


def report_text(meta: dict, findings: List[Finding], output: str):
    findings = sort_findings(findings)
    lines = ["=" * 78, " MOBILE APP SECURITY SCAN REPORT", "=" * 78,
             f"Target  : {meta['target']}", f"Type    : {meta['type']}",
             f"SHA256  : {meta['sha256']}", f"Tools   : {meta['tooling']}",
             f"Findings: {len(findings)}"]
    counts = build_summary(findings)
    lines.append("Severity: " + ", ".join(f"{k}={v}" for k, v in counts.items() if v))
    fp = sum(1 for f in findings if _conf_val(f) == Confidence.LIKELY_FALSE_POSITIVE.value)
    lines.append(f"Likely false positives (auto-downgraded): {fp}")
    lines.append("")
    for i, f in enumerate(findings, 1):
        flag = " [LIKELY FALSE POSITIVE]" if _conf_val(f) == Confidence.LIKELY_FALSE_POSITIVE.value else ""
        lines += ["-" * 78,
                  f"{i}. [{_as_sev(f.severity).value}]{flag} {f.title}  ({f.check_id})",
                  f"   Confidence : {_conf_val(f)}",
                  f"   Location   : {f.location}",
                  f"   Evidence   : {f.evidence}",
                  f"   Description: {f.description}",
                  f"   Validation : {f.validation}"]
        if f.remediation:
            lines.append(f"   Remediation: {f.remediation}")
        lines.append("")
    with open(output, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    print(f"Text report written to {output}")


def report_json(meta: dict, findings: List[Finding], output: str):
    findings = sort_findings(findings)
    report = {**meta, "summary": build_summary(findings),
              "findings": [f.to_dict() for f in findings]}
    with open(output, "w", encoding="utf-8") as fh:
        json.dump(report, fh, indent=2)
    print(f"JSON report written to {output}")


def report_word(meta: dict, findings: List[Finding], output: str):
    if not HAVE_DOCX:
        print("error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(2)
    sev_colors = {
        Severity.CRITICAL: RGBColor(0x8B, 0x00, 0x00), Severity.HIGH: RGBColor(0xC0, 0x00, 0x00),
        Severity.MEDIUM: RGBColor(0xE0, 0x7A, 0x00), Severity.LOW: RGBColor(0x1F, 0x6F, 0x1F),
        Severity.INFO: RGBColor(0x40, 0x40, 0x40),
    }
    findings = sort_findings(findings)
    doc = Document()
    t = doc.add_heading("Mobile App Security Scan Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Scan Metadata", level=1)
    mt = doc.add_table(rows=0, cols=2); mt.style = "Light Grid Accent 1"
    for k, v in [("Target", meta["target"]), ("Type", meta["type"]), ("SHA256", meta["sha256"]),
                 ("Tooling", meta["tooling"]), ("Total Findings", str(len(findings)))]:
        c = mt.add_row().cells; c[0].text = k; c[1].text = str(v)
    doc.add_heading("Severity Summary", level=1)
    counts = build_summary(findings)
    st = doc.add_table(rows=1, cols=2); st.style = "Light Grid Accent 1"
    h = st.rows[0].cells; h[0].text, h[1].text = "Severity", "Count"
    for sev in [Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO]:
        r = st.add_row().cells; r[0].text = sev.value; r[1].text = str(counts[sev.value])
    fp = sum(1 for f in findings if _conf_val(f) == Confidence.LIKELY_FALSE_POSITIVE.value)
    doc.add_paragraph().add_run(f"Likely false positives (auto-downgraded): {fp}").italic = True
    doc.add_heading("Detailed Findings", level=1)
    for i, f in enumerate(findings, 1):
        hd = doc.add_heading(level=2)
        rr = hd.add_run(f"{i}. [{_as_sev(f.severity).value}] {f.title}")
        rr.font.color.rgb = sev_colors[_as_sev(f.severity)]
        if _conf_val(f) == Confidence.LIKELY_FALSE_POSITIVE.value:
            hd.add_run("  (LIKELY FALSE POSITIVE)").italic = True
        det = doc.add_table(rows=0, cols=2); det.style = "Light List Accent 1"
        for label, value in [("Check ID", f.check_id), ("Severity", _as_sev(f.severity).value),
                             ("Confidence", _conf_val(f)), ("Location", f.location),
                             ("Evidence", f.evidence), ("Description", f.description),
                             ("Validation", f.validation), ("Remediation", f.remediation or "N/A")]:
            cells = det.add_row().cells
            cells[0].text = label; cells[1].text = str(value)
            for para in cells[0].paragraphs:
                for run_ in para.runs:
                    run_.bold = True
        doc.add_paragraph()
    doc.save(output)
    print(f"Word report written to {output}")


# --------------------------------------------------------------------------- #
# Orchestration
# --------------------------------------------------------------------------- #
def detect_type(path: str) -> str:
    lower = path.lower()
    if lower.endswith(".apk"):
        return "APK"
    if lower.endswith(".ipa"):
        return "IPA"
    try:
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
            if any(n == "AndroidManifest.xml" or n.startswith("classes") for n in names):
                return "APK"
            if any(n.startswith("Payload/") for n in names):
                return "IPA"
    except Exception:
        pass
    return "UNKNOWN"


def tooling_summary(decompile: bool, cve: bool, dynamic: bool) -> str:
    bits = ["androguard" if HAVE_ANDROGUARD else "androguard:MISSING",
            "macholib" if HAVE_MACHOLIB else "macholib:MISSING",
            "python-docx" if HAVE_DOCX else "python-docx:MISSING",
            "frida" if HAVE_FRIDA else "frida:MISSING",
            "jadx" if have_tool("jadx") else "jadx:MISSING",
            "apktool" if have_tool("apktool") else "apktool:MISSING"]
    if cve:
        bits.append("trivy" if have_tool("trivy") else ("grype" if have_tool("grype") else "cve:NONE"))
    if not decompile:
        bits.append("decompile:OFF")
    bits.append("dynamic:ON" if dynamic else "dynamic:OFF")
    return ", ".join(bits)


def run_validation(findings: List[Finding], engine: ValidationEngine) -> List[Finding]:
    seen_ids = {f.check_id for f in findings}
    adjusted = []
    for f in findings:
        ctx = FindingContext(
            file_path=f.location, package=f.location, raw_match=f.evidence,
            quoted_value=_extract_quoted(f.evidence),
            corroborating_ids=list(seen_ids - {f.check_id}),
            method_signatures=_methods_from_finding(f),
            build_variant=_guess_build_variant(f.location),
            runtime_observed=(f.location == "[runtime]"),
        )
        result = engine.evaluate(f, ctx)
        if result is not None:
            adjusted.append(result)
    engine.finalize()
    return adjusted


def main():
    ap = argparse.ArgumentParser(description="APK/IPA vulnerability scanner v4 (static + dynamic).")
    ap.add_argument("target", help="Path to .apk or .ipa file")
    ap.add_argument("--format", choices=["console", "text", "json", "word"],
                    default="console", help="Output format (default: console)")
    ap.add_argument("--output", help="Output file (required for text/json/word)")
    ap.add_argument("--min-severity", default="INFO", choices=[s.value for s in Severity])
    ap.add_argument("--no-decompile", action="store_true", help="Skip jadx/apktool")
    ap.add_argument("--no-cve", action="store_true", help="Skip trivy/grype")
    ap.add_argument("--baseline", help="Baseline JSON for FP suppression across runs")
    ap.add_argument("--auto-suppress", action="store_true",
                    help="Auto-add strongly-exonerated findings to the baseline")
    ap.add_argument("--keep-workdir", help="Keep extracted/decompiled output here")
    # dynamic
    ap.add_argument("--dynamic", action="store_true", help="Run Frida dynamic analysis")
    ap.add_argument("--dyn-target", help="package id (Android) / app name (iOS) for dynamic run")
    ap.add_argument("--dyn-host", help="remote frida host (e.g. Corellium device IP)")
    ap.add_argument("--dyn-attach", action="store_true", help="attach instead of spawn")
    ap.add_argument("--dyn-duration", type=int, default=30, help="dynamic instrumentation seconds")
    args = ap.parse_args()

    if not os.path.isfile(args.target):
        print(f"error: file not found: {args.target}", file=sys.stderr)
        sys.exit(1)
    if args.format in ("text", "json", "word") and not args.output:
        print(f"error: --output is required for --format {args.format}", file=sys.stderr)
        sys.exit(1)
    if args.dynamic and not args.dyn_target:
        print("error: --dynamic requires --dyn-target", file=sys.stderr)
        sys.exit(1)

    ftype = detect_type(args.target)
    if ftype == "UNKNOWN":
        print("error: not a valid APK/IPA.", file=sys.stderr)
        sys.exit(1)

    decompile = not args.no_decompile
    do_cve = not args.no_cve
    digest = sha256_of(args.target)
    findings: List[Finding] = []

    workdir = args.keep_workdir or tempfile.mkdtemp(prefix="mobscan_")
    cleanup = not args.keep_workdir
    try:
        if ftype == "APK":
            core, source_roots = analyze_apk(args.target, workdir, decompile)
        else:
            core, source_roots = analyze_ipa(args.target, workdir)
        findings += core
        findings += scan_secrets(source_roots)
        findings += scan_dependencies(workdir, do_cve)
    finally:
        if cleanup:
            shutil.rmtree(workdir, ignore_errors=True)
        else:
            print(f"[workdir preserved at: {workdir}]")

    # ---- Dynamic phase ----
    if args.dynamic:
        plat = "android" if ftype == "APK" else "ios"
        raw_dyn = run_dynamic(plat, args.dyn_target, host=args.dyn_host,
                              spawn=not args.dyn_attach, duration=args.dyn_duration)
        findings += dynamic_to_findings(raw_dyn)
        print(f"[DYN] Collected {len(raw_dyn)} dynamic finding(s).")

    # ---- Validation engine ----
    engine = ValidationEngine(baseline_path=args.baseline, auto_suppress=args.auto_suppress)
    findings = run_validation(findings, engine)

    # ---- Severity filter ----
    threshold = SEV_ORDER[Severity(args.min_severity)]
    findings = [f for f in findings if SEV_ORDER[_as_sev(f.severity)] <= threshold]

    meta = {"target": args.target, "type": ftype, "sha256": digest,
            "tooling": tooling_summary(decompile, do_cve, args.dynamic)}

    if args.format == "console":
        report_console(meta, findings)
    elif args.format == "text":
        report_text(meta, findings, args.output)
    elif args.format == "json":
        report_json(meta, findings, args.output)
    elif args.format == "word":
        report_word(meta, findings, args.output)


if __name__ == "__main__":
    main()